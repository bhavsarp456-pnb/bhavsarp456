import docx

def sum_num(a,b):
    return a+b
res =sum_num(15,5)


doc = docx.Document()
doc.add_heading("Sum of two numbers", 0)
Sum_para = doc.add_paragraph("Sum of the Function is \t")
Sum_para.add_run(f"{res}").bold = True
doc.add_page_break()

doc.save("C:\\Users\\Dell\\OneDrive\\Desktop\\Sum1.docx")