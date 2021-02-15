import docx
doc = docx.Document()
doc.add_heading("RST Forum", 0)
network_para = doc.add_paragraph("Newtorking")
py_para = doc.add_paragraph("Python")
cloud_para = doc.add_paragraph("Cloud")
sec_para = doc.add_paragraph("Security")
col_para = doc.add_paragraph("Collaboration")
serv_para = doc.add_paragraph("Servers")

network_para.add_run("Foundation of next generationtechnology").bold = True
py_para.add_run("Foundation of next Automation").italic = True
cloud_para.add_run("Reduce capex & opex").underline = True
sec_para.add_run("Name hai Piyush, raheta hu binnddasss")
col_para.add_run("Converged architecture")
serv_para.add_run("Goli nahi marenge, keh ke lenge uski...")

doc.add_page_break()
doc.add_heading("Coming up Next Gen Tech", 2)
doc.add_picture("C:\\Users\\Dell\\OneDrive\\Desktop\\PiyushProfilePic.jpg")

doc.save("C:\\Users\\Dell\\OneDrive\\Desktop\\pydoc.docx")